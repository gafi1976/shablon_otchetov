# generator_ust.py
# Генератор Word актов УСТАНОВКИ по шаблону shablom_ust.docx
#
# Структура выходного документа (один файл = один акт / один набор оборудования):
#
# «TASDIQLAYMAN» {Organizasiya} rahbari _______ {boss name}
# «{day}» {month} {year} yil
#
# {Organizasiya}
# Manzil: {adress}
#
# USKUNANI O'RNATISH DALOLATNOMASI
#
# Dalolatnoma № {num_otch}
#
# Mazkur dalolatnoma ... maqsadida tuzildi.
#
# [ТАБЛИЦА]:
#   № | O'rnatilgan sana | Uskuna | Seriya raqami | Qayerga o'rnatilgan
#   {num} | {date} | {oborud name} | {serial num} | {where oborud}
#
# Yuqorida ko'rsatilgan uskunalar ... ekspluatatsiyaga topshirildi.
#
# {job title1}: ____  {engine1}
# {job title2}: ____  {engine2}

import os
import copy
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
#  Вспомогательные функции
# ─────────────────────────────────────────────

MONTHS_UZ = {
    1:'yanvar', 2:'fevral', 3:'mart', 4:'aprel',
    5:'may', 6:'iyun', 7:'iyul', 8:'avgust',
    9:'sentabr', 10:'oktabr', 11:'noyabr', 12:'dekabr'
}

def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _set_cell_border(cell, color='1A5276'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=1.5):
    sec = doc.sections[0]
    sec.top_margin    = Cm(top)
    sec.bottom_margin = Cm(bottom)
    sec.left_margin   = Cm(left)
    sec.right_margin  = Cm(right)

def _p(doc, text='', bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
       color=None, before=0, after=4, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = Pt(15)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def _hr(doc, color='1A5276'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    e = OxmlElement('w:bottom')
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), '12')
    e.set(qn('w:space'), '1')
    e.set(qn('w:color'), color)
    pBdr.append(e)
    pPr.append(pBdr)


# ─────────────────────────────────────────────
#  Основной генератор — один документ на
#  один акт установки (один набор оборудования)
# ─────────────────────────────────────────────

def create_ust_doc(data: dict, doc_number: str = '1') -> Document:
    """
    data = {
        'org_name':        str,   # {Organizasiya}
        'region':          str,   # адрес в заголовке
        'address':         str,   # Manzil: {adress}
        'commission_head': str,   # {boss name}
        'member1':         str,   # {engine1}
        'member1_title':   str,   # {job title1}
        'member2':         str,   # {engine2}
        'member2_title':   str,   # {job title2}
        'doc_number':      str,   # {num_otch}
        'doc_date':        str,   # дд.мм.гггг
        'items': [
            {
                'num':           str,
                'name':          str,   # {oborud name}
                'serial_number': str,   # {serial num}
                'install_date':  str,   # {installation date}
                'location':      str,   # {where oborud}
                'inv_number':    str,
                'model':         str,
                'cost':          str,
                'condition':     str,
                'note':          str,
            }
        ]
    }
    """
    doc = Document()
    _set_margins(doc)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    today    = datetime.now()
    date_str = data.get('doc_date', today.strftime('%d.%m.%Y'))
    num_otch = data.get('doc_number', doc_number)
    try:
        dt    = datetime.strptime(date_str, '%d.%m.%Y')
        day   = str(dt.day)
        month = MONTHS_UZ[dt.month]
        year  = str(dt.year)
    except Exception:
        day   = str(today.day)
        month = MONTHS_UZ[today.month]
        year  = str(today.year)

    org_name  = data.get('org_name', '')
    address   = data.get('address', data.get('region', ''))
    boss      = data.get('commission_head', '')
    eng1      = data.get('member1', '')
    job1      = data.get('member1_title', 'Yetakchi muhandis')
    eng2      = data.get('member2', '')
    job2      = data.get('member2_title', 'Yetakchi muhandis')
    items     = data.get('items', [])

    # ══════════════════════════════════
    #  ШАПКА — «TASDIQLAYMAN»
    # ══════════════════════════════════
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after  = Pt(0)

    def rr(p, text, bold=False, size=11):
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = 'Times New Roman'

    rr(p_hdr, '«TASDIQLAYMAN»\n', bold=True, size=11)
    rr(p_hdr, f'{org_name} rahbari\n', size=10)
    rr(p_hdr, f'_______________________________________ {boss}\n', size=10)
    rr(p_hdr, f'«{day}» {month} {year} yil', size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _hr(doc)

    # ══════════════════════════════════
    #  ОРГАНИЗАЦИЯ И АДРЕС
    # ══════════════════════════════════
    _p(doc, org_name,
       bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER,
       color=(0x1A, 0x52, 0x76), before=6, after=2)

    _p(doc, f"Manzil: {address}",
       size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
       italic=True, after=8, color=(0x55, 0x55, 0x55))

    # ══════════════════════════════════
    #  ЗАГОЛОВОК ДALOLATNOMA
    # ══════════════════════════════════
    _p(doc, "USKUNANI O\u2019RNATISH DALOLATNOMASI",
       bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER,
       color=(0x1A, 0x52, 0x76), before=4, after=3)

    _p(doc, f"Dalolatnoma \u2116{num_otch}",
       bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
       before=0, after=8)

    # ══════════════════════════════════
    #  ВВОДНЫЙ ТЕКСТ
    # ══════════════════════════════════
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.space_before = Pt(0)
    intro.paragraph_format.space_after  = Pt(8)
    intro.paragraph_format.line_spacing = Pt(16)
    ir = intro.add_run(
        "\tMazkur dalolatnoma texnologik jarayonlarni modernizatsiya qilish "
        "doirasida quyidagi uskunalar o\u2019rnatilganligini tasdiqlash "
        "maqsadida tuzildi."
    )
    ir.font.size = Pt(12)
    ir.font.name = 'Times New Roman'

    # ══════════════════════════════════
    #  ТАБЛИЦА ОБОРУДОВАНИЯ
    # ══════════════════════════════════
    # № | O'rnatilgan sana | Uskuna | Seriya raqami | Qayerga o'rnatilgan
    COL_WIDTHS   = [Cm(1.2), Cm(2.8), Cm(5.0), Cm(3.5), Cm(5.0)]
    HEADER_COLOR = '1A5276'
    EVEN_BG      = 'EBF5FB'
    ODD_BG       = 'FFFFFF'

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Заголовок таблицы
    hdr_labels = [
        '\u2116',
        "O\u2019rnatilgan\nsana",
        'Uskuna\n(Model)',
        'Seriya\nraqami',
        "Qayerga\no\u2019rnatilgan",
    ]
    hdr_row = table.rows[0]
    hdr_row.height = Pt(28)
    for ci, (cell, label, w) in enumerate(zip(hdr_row.cells, hdr_labels, COL_WIDTHS)):
        cell.width = w
        _set_cell_bg(cell, HEADER_COLOR)
        _set_cell_border(cell, HEADER_COLOR)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Строки данных
    for i, item in enumerate(items):
        row = table.add_row()
        row.height = Pt(22)
        bg = EVEN_BG if i % 2 == 0 else ODD_BG

        inst_date = item.get('install_date', date_str) or date_str
        values = [
            item.get('num', str(i + 1)),
            inst_date,
            item.get('name', item.get('model', '')),
            item.get('serial_number', ''),
            item.get('location', item.get('note', '')),
        ]

        for ci, (cell, val) in enumerate(zip(row.cells, values)):
            _set_cell_bg(cell, bg)
            _set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.line_spacing = Pt(12)
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ══════════════════════════════════
    #  ЗАКЛЮЧИТЕЛЬНЫЙ ТЕКСТ
    # ══════════════════════════════════
    outro = doc.add_paragraph()
    outro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    outro.paragraph_format.space_before = Pt(0)
    outro.paragraph_format.space_after  = Pt(10)
    outro.paragraph_format.line_spacing = Pt(16)
    or_ = outro.add_run(
        "\tYuqorida ko\u2019rsatilgan uskunalar belgilangan tartibda "
        "o\u2019rnatildi, tekshirildi va ekspluatatsiyaga topshirildi."
    )
    or_.font.size = Pt(12)
    or_.font.name = 'Times New Roman'

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ══════════════════════════════════
    #  ПОДПИСИ
    # ══════════════════════════════════
    sig_tbl = doc.add_table(rows=2, cols=3)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (job, engineer) in enumerate([
        (job1, eng1),
        (job2, eng2),
    ]):
        row = sig_tbl.rows[i]
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(6)
        row.cells[2].width = Cm(5)

        def sc(cell, text, bold=False):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(8)
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(11)
            r.font.name = 'Times New Roman'

        sc(row.cells[0], f"O\u2019rnatish uchun mas\u2019ul shaxs\n{job}:", bold=True)
        sc(row.cells[1], '________________________')
        sc(row.cells[2], engineer)

    # ── Нижняя линия ──
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p_foot._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top_e = OxmlElement('w:top')
    top_e.set(qn('w:val'), 'single')
    top_e.set(qn('w:sz'), '12')
    top_e.set(qn('w:space'), '1')
    top_e.set(qn('w:color'), '1A5276')
    pBdr.append(top_e)
    pPr.append(pBdr)
    fr = p_foot.add_run(
        f"М.П.   Dalolatnoma \u2116{num_otch}   |   Sana: {date_str}"
    )
    fr.font.size = Pt(9)
    fr.font.name = 'Times New Roman'
    fr.italic = True
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


# ─────────────────────────────────────────────
#  Сохранение файлов
# ─────────────────────────────────────────────

def save_single_files(data_list: list, output_dir: str) -> list:
    """
    Один .docx файл на каждую строку оборудования (каждый item).
    Имя файла: Dalolatnoma_{num}_{oborud_name}.docx
    """
    os.makedirs(output_dir, exist_ok=True)
    created = []

    for data in data_list:
        items = data.get('items', [])
        # Если items пустой — один файл на весь data
        if not items:
            doc  = create_ust_doc(data, doc_number='1')
            name = f"Dalolatnoma_1.docx"
            path = os.path.join(output_dir, name)
            doc.save(path)
            created.append(path)
            continue

        # Каждый item = отдельный файл
        for item in items:
            single_data = {**data, 'items': [item]}
            num  = str(item.get('num', '')).strip()
            name_raw = item.get('name', item.get('model', '')).strip()
            # Убираем недопустимые символы для имени файла
            safe_name = ''.join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in name_raw)
            safe_name = safe_name.strip()[:40]
            fname = f"Dalolatnoma_{num}_{safe_name}.docx" if safe_name else f"Dalolatnoma_{num}.docx"
            path  = os.path.join(output_dir, fname)
            doc   = create_ust_doc(single_data, doc_number=num or str(len(created)+1))
            doc.save(path)
            created.append(path)

    return created


def save_all_in_one(data_list: list, output_path: str) -> str:
    """Все акты в одном файле, разделённые разрывом страницы."""
    combined = Document()
    combined.sections[0].top_margin    = Cm(2.0)
    combined.sections[0].bottom_margin = Cm(2.0)
    combined.sections[0].left_margin   = Cm(2.5)
    combined.sections[0].right_margin  = Cm(1.5)
    combined.styles['Normal'].font.name = 'Times New Roman'
    combined.styles['Normal'].font.size = Pt(12)

    for idx, data in enumerate(data_list, start=1):
        single = create_ust_doc(data, doc_number=str(idx))
        for element in single.element.body:
            combined.element.body.append(copy.deepcopy(element))
        if idx < len(data_list):
            pb = OxmlElement('w:p')
            pb_r = OxmlElement('w:r')
            pb_br = OxmlElement('w:br')
            pb_br.set(qn('w:type'), 'page')
            pb_r.append(pb_br)
            pb.append(pb_r)
            combined.element.body.append(pb)

    combined.save(output_path)
    return output_path
