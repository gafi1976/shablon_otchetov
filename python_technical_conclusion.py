import pandas as pd
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
import os

# ========== 1. ЗАГРУЗКА ДАННЫХ ИЗ EXCEL И ОЧИСТКА ОТ ИТОГОВЫХ СТРОК ==========
file_path = "data.xlsx"

# Читаем все листы Excel файла
excel_file = pd.ExcelFile(file_path)
all_sheets = {}

for sheet_name in excel_file.sheet_names:
    df = pd.read_excel(file_path, sheet_name=sheet_name)

    # Проверка колонок для каждого листа
    required_columns = ['Qurilma nomi', 'Qismlar nomi', 'Foydalanishga yaroqliligi', 'Nosozlik belgilari',
                        'Inventar raqami']
    missing_cols = [col for col in required_columns if col not in df.columns]

    if missing_cols:
        print(f"⚠️  Лист '{sheet_name}' пропущен (нет колонок: {missing_cols})")
        continue

    # УДАЛЯЕМ строки, где в колонке 'Қурилма номи' есть "ИТОГОВОЕ РЕШЕНИЕ"
    before_count = len(df)
    df = df[~df['Qurilma nomi'].astype(str).str.contains('ИТОГОВОЕ РЕШЕНИЕ', na=False)]
    after_count = len(df)

    if before_count != after_count:
        print(f"   🗑️  Лист '{sheet_name}': удалено {before_count - after_count} итоговых строк")

    all_sheets[sheet_name] = df
    print(f"✅ Лист '{sheet_name}': загружено {len(df)} записей")

if not all_sheets:
    raise ValueError("Нет ни одного листа с правильными колонками!")

# ========== 2. ПАРАМЕТРЫ ==========
today = datetime.now()
region = "Sirdaryo viloyati"
organization = "ABM MChJ"
signer = "A.C. Babayev"
engineer1 = "V.A. Raxmatov"
engineer2 = "G'.T. Xolbekov"


# Функция для получения следующего номера заключения
def get_next_doc_number():
    """Автоматически определяет следующий номер заключения"""
    existing_files = [f for f in os.listdir('.') if f.startswith('technical_conclusion_') and f.endswith('.docx')]

    if not existing_files:
        return 1

    numbers = []
    for file in existing_files:
        # Извлекаем номер из имени файла (если есть)
        import re
        match = re.search(r'_(\d+)_', file)
        if match:
            numbers.append(int(match.group(1)))
        else:
            # Если в имени нет номера, проверяем внутри файла (медленнее)
            # Для простоты используем количество файлов + 1
            pass

    if numbers:
        return max(numbers) + 1
    else:
        return len(existing_files) + 1


# ========== 3. ФУНКЦИЯ СОЗДАНИЯ WORD ДОКУМЕНТА ==========
def create_word_conclusion_for_device(dataframe, sheet_name, device_name, doc_number):
    """Создает Word заключение для одного оборудования"""
    doc = Document()

    # Настройка стиля и интервалов
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1

    # ШАПКА (ТАСДИҚЛАЙМАН)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run("«TASDIQLAYMAN»\n").bold = True
    p.add_run(f"{organization} {region}\n")
    p.add_run(f"bo'lim boshlig'i\n")
    p.add_run(f"__________ {signer}\n\n")
    p.add_run(f"{today.strftime('%Y yil %d %B')}\n")

    doc.add_paragraph()

    # ЗАГОЛОВОК (с автоматическим номером)
    title = doc.add_heading(f"TEXNIK XULOSA № {doc_number}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(6)

    # ВВОДНАЯ ЧАСТЬ
    p = doc.add_paragraph(
        f"\tBiz, quyida imzo chekuvchilar — {organization} {region} bo'lim yetakchi muhandisi {engineer1} "
        f"va yetakchi muhandis {engineer2}lar, quyida keltirilgan qurilmalarni texnik ko'rikdan o'tkazdik:"
    )
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

    # ИНФОРМАЦИЯ ОБ ОБОРУДОВАНИИ (перед таблицей)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Qurilma nomi: ").bold = True
    p.add_run(f"{device_name}")

    # Получаем инвентарный номер из первой строки (он одинаковый для всех компонентов)
    inv_number = dataframe.iloc[0]['Inventar raqami'] if len(dataframe) > 0 else "Н/Д"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Inventar raqami: ").bold = True
    p.add_run(f"{inv_number}")

    doc.add_paragraph()

    # ТАБЛИЦА С ДАННЫМИ (только 3 колонки: Қисмлар номи, Фойдаланишга яроқлилиги, Носозлик белгилари)
    num_rows = len(dataframe) + 1
    table = doc.add_table(rows=num_rows, cols=3)
    table.style = 'Table Grid'

    # Заголовки таблицы (3 колонки)
    headers = ['Qismlar nomi', 'Foydalanishga yaroqliligi', 'Nosozlik belgilari']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.runs[0].font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    # Заполнение данными
    for row_idx, (_, row) in enumerate(dataframe.iterrows(), start=1):
        cells_data = [
            str(row['Qismlar nomi']),
            str(row['Foydalanishga yaroqliligi']),
            str(row['Nosozlik belgilari'])
        ]

        is_defective = row['Foydalanishga yaroqliligi'] == 'Yaroqsiz'

        for col_idx, cell_data in enumerate(cells_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_data
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

            if col_idx == 1 and is_defective:  # колонка "Фойдаланишга яроқлилиги"
                for paragraph in cell.paragraphs:
                    paragraph.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    paragraph.runs[0].font.bold = True

    # Отступ после таблицы
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # ХУЛОСА (по центру)
    heading = doc.add_heading("XULOSA", level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)

    # Статистика для определения текста заключения
    defective_parts = len(dataframe[dataframe['Foydalanishga yaroqliligi'] == 'Yaroqsiz'])
    moral_obsolete = len(
        dataframe[dataframe["Nosozlik belgilari"].str.contains("ma'naviy eskirgan|eskirgan", case=False, na=False)])

    # Текст заключения
    if moral_obsolete > 0 or defective_parts > 0:
        p = doc.add_paragraph(f"\t«{device_name}» zamonaviy talablarga javob bermaydi, ma'naviy eskirgan.")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph(
            f"Ta'mirlash uchun zarur bo'lgan ehtiyot qismlar texnologik jarayondan olib tashlanganligini hisobga olib, "
            f"uni tiklash maqsadga muvofiq emas."
        )
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        doc.add_paragraph()

        p = doc.add_paragraph(
            f"\tYuqorida keltirib o'tilganlardan kelib chiqib, komissiya «{device_name}» asosiy vositalar "
            f"ro'yxatidan chiqarilsin degan xulosaga keldi."
        )
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    else:
        p = doc.add_paragraph(f"«{device_name}» tekshiruvdan so'ng foydalanishga yaroqli deb topildi.")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Отступ перед подписями
    doc.add_paragraph()
    doc.add_paragraph()

    # ПОДПИСИ
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run("Yetakchi muhandis:").bold = True
    p.add_run("                                                            ")
    p.add_run("______________________").bold = True
    p.add_run(f"  {engineer1}")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.add_run("Yetakchi muhandis:").bold = True
    p.add_run("                                                            ")
    p.add_run("______________________").bold = True
    p.add_run(f"  {engineer2}")

    return doc


# ========== 4. ФУНКЦИЯ СОХРАНЕНИЯ EXCEL (БЕЗ ИТОГОВЫХ СТРОК) ==========
def save_clean_excel(all_data, filename="data.xlsx"):
    """Сохраняет Excel без итоговых строк, только с исходными данными"""

    with pd.ExcelWriter(filename, engine='openpyxl', mode='w') as writer:
        for sheet_name, dataframe in all_data.items():
            # Сохраняем очищенные данные
            dataframe.to_excel(writer, sheet_name=sheet_name, index=False)

            # Получаем workbook и worksheet для форматирования
            workbook = writer.book
            worksheet = writer.sheets[sheet_name]

            # Настройка стилей
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            red_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
            red_font = Font(color="9C0006", bold=True)

            # Форматирование заголовков
            for col in range(1, len(dataframe.columns) + 1):
                cell = worksheet.cell(row=1, column=col)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')

            # Форматирование данных - подсветка строк с Яроқсиз
            for row in range(2, len(dataframe) + 2):
                condition_cell = worksheet.cell(row=row, column=3)

                # Подсветка красным если "Яроқсиз"
                if condition_cell.value == "Яроқсиз":
                    for col in range(1, 6):
                        cell = worksheet.cell(row=row, column=col)
                        cell.fill = red_fill
                        cell.font = red_font

            # Автоширина колонок
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 30)
                worksheet.column_dimensions[column_letter].width = adjusted_width

    print(f"✅ Excel файл сохранён (без итоговых строк): {filename}")


# ========== 5. ГЛАВНАЯ ФУНКЦИЯ ==========
def main():
    # Получаем следующий номер заключения
    start_number = get_next_doc_number()
    print(f"📄 Начальный номер заключения: {start_number}")

    # Создаем Word заключения для каждого оборудования
    doc_number = start_number
    for sheet_name, df_sheet in all_sheets.items():
        # Получаем название оборудования из первого ряда первой колонки
        device_name = df_sheet.iloc[0]['Qurilma nomi']

        # Создаем Word документ
        doc = create_word_conclusion_for_device(df_sheet, sheet_name, device_name, doc_number)

        # Сохраняем с именем, содержащим номер
        safe_name = device_name.replace(' ', '_').replace('/', '_')
        word_file = f"texnik_xulosa_{doc_number}_{safe_name}.docx"
        doc.save(word_file)
        print(f"✅ Word заключение №{doc_number} для '{device_name}': {word_file}")

        doc_number += 1

    # Сохраняем Excel без итоговых строк
    save_clean_excel(all_sheets, "data.xlsx")

    # Выводим общую статистику
    print("\n" + "=" * 60)
    print("ОБЩАЯ СТАТИСТИКА ПО ВСЕМУ ОБОРУДОВАНИЮ")
    print("=" * 60)

    total_devices = len(all_sheets)
    total_components = 0
    total_defective = 0

    for sheet_name, df_sheet in all_sheets.items():
        device_name = df_sheet.iloc[0]['Qurilma nomi']
        components = len(df_sheet)
        defective = len(df_sheet[df_sheet['Foydalanishga yaroqliligi'] == 'Yaroqsiz'])

        total_components += components
        total_defective += defective

        print(f"\n📌 {device_name}:")
        print(f"   - Компонентов: {components}")
        print(f"   - Неисправных: {defective}")

    print(f"\n📊 ИТОГО:")
    print(f"   - Всего оборудований: {total_devices}")
    print(f"   - Всего компонентов: {total_components}")
    print(f"   - Всего неисправных: {total_defective}")
    print(f"   - Номера заключений: {start_number} - {doc_number - 1}")
    print("=" * 60)

    print("\n📁 Созданные файлы:")
    print("   • texnik_xulosa_1_*.docx - заключения для каждого оборудования")
    print("   • data.xlsx - очищенный Excel (без итоговых строк, только данные)")


if __name__ == "__main__":
    main()