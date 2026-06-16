# generator_spisan.py
# Альбомный акт списания основных средств
# Образец: AKTOC2026N20.docx

import os
import copy
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from transliterate import TEXTS_SPISAN, MONTHS, convert_to, detect_script



# ─────────────────────────────────────────────
#  Вспомогательные функции
# ─────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color='333333'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _no_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _p(doc, text='', bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
       color=None, before=0, after=2, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = Pt(12)
    if text:
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.size = Pt(size); run.font.name = 'Times New Roman'
        if color: run.font.color.rgb = RGBColor(*color)
    return p


def _cell_p(cell, text='', bold=False, size=9,
            align=WD_ALIGN_PARAGRAPH.CENTER, italic=False, color=None):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    if text:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = 'Times New Roman'
        if color: r.font.color.rgb = RGBColor(*color)
    return p


def _auto_lang(group: dict) -> str:
    for key in ('org_name', 'commission_head', 'member1', 'member2'):
        val = group.get(key, '').strip()
        if val:
            s = detect_script(val)
            if s in ('latin', 'cyrillic'): return s
    for dev in group.get('devices', []):
        s = detect_script(dev.get('name', ''))
        if s in ('latin', 'cyrillic'): return s
    return 'latin'



def _set_landscape(doc):
    """Альбомная ориентация А4, минимальные поля как в образце."""
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin    = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    # Устанавливаем альбомную ориентацию через XML
    sectPr = section._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.append(pgSz)
    pgSz.set(qn('w:orient'), 'landscape')


# ─────────────────────────────────────────────
#  ТЕКСТЫ документа (кириллица — образец)
# ─────────────────────────────────────────────

TEXTS_CYR = {
    'tasdiq':     'Тасдиқлайман:',
    'title1':     'АСОСИЙ ВОСИТАЛАРНИ РЎЙХАТДАН ЧИҚАРИШ БЎЙИЧА',
    'title2':     'ДАЛОЛАТНОМА',
    'komissiya':  'АБМ {org} томонидан {date} йил {order} сонли буйруғи билан тайинланган комиссия қуйидагиларга асосан',
    'ish_chiq':   'йилда ишлаб чиқилган.',
    'kelgan':     'йилда корхонага келган.',
    'ishga':      'йилда ишга туширилган.',
    'kapital':    'Капитал таъмирлаш сони _____ унга _________ сўм харажат қилинган.',
    'texnik':     'Техник ҳолати:   маънавий эскирган, ишга яроқсиз ҳолга келган, тузатиш мақсадга мувофиқ эмас.',
    'xulosa_txt': 'Комиссия хулосаси: {org} да ишлатилган {names} {date} йил «{day}» {month} "Техник хулоса" га асосан асосий воситалар рўйхатидан чиқарилсин.',
    'moddiy':     'Бўлим бўйича моддий жавобгар шахс:',
    'komrais':    'Комиссия раиси {org}:',
    'komazolari': 'Комиссия аъзолари:',
    'muhandis':   'Етакчи муҳандис',
    'rahbar_title': 'раҳбари',
    # Таблица
    'bulim':      'Бўлим',
    'debet':      'Дебет\nҲисобварақ\nкоди',
    'kredit':     'Кредит\nҲисобварақ\nкоди',
    'nomi':       'Номи',
    'narxi':      'Нархи\n(сўм)',
    'inv':        'Инв.№',
    'amort':      'Амортизация\nчегирмалари\nсуммаси',
    'norma':      'Амортизация\nнормаси\n(%)',
    'qoldiq':     'Қолдиқ\nнархи',
}

TEXTS_LAT = {
    'tasdiq':     "Tasdiqlayman:",
    'title1':     "ASOSIY VOSITALARNI RO'YXATDAN CHIQARISH BO'YICHA",
    'title2':     'DALOLATNOMA',
    'komissiya':  "ABM {org} tomonidan {date} yil {order} sonli buyrug'i bilan tayinlangan komissiya quyidagilarga asosan",
    'ish_chiq':   'yilda ishlab chiqilgan.',
    'kelgan':     'yilda korxonaga kelgan.',
    'ishga':      'yilda ishga tushirilgan.',
    'kapital':    "Kapital ta'mirlash soni _____ unga _________ so'm xarajat qilingan.",
    'texnik':     "Texnik holati: ma'naviy eskirgan, ishga yaroqsiz holga kelgan, tuzatish maqsadga muvofiq emas.",
    'xulosa_txt': "Komissiya xulosasi: {org} da ishlatilingan {names} {date} yil «{day}» {month} \"Texnik xulosa\" ga asosan asosiy vositalar ro'yxatidan chiqarilsin.",
    'moddiy':     "Bo'lim bo'yicha moddiy javobgar shaxs:",
    'komrais':    "Komissiya raisi {org}:",
    'komazolari': "Komissiya a'zolari:",
    'muhandis':   'Yetakchi muhandis',
    'rahbar_title': 'rahbari',
    'bulim':      "Bo'lim",
    'debet':      'Debet\nHisob\nraqam',
    'kredit':     'Kredit\nHisob\nraqam',
    'nomi':       'Nomi',
    'narxi':      "Narxi\n(so'm)",
    'inv':        'Inv.№',
    'amort':      'Amortizatsiya\negirmalari\nsummasi',
    'norma':      'Amortizatsiya\nnormasi\n(%)',
    'qoldiq':     'Qoldiq\nnarxi',
}



# ─────────────────────────────────────────────
#  Основной генератор — один документ
# ─────────────────────────────────────────────

def create_spisan_doc(group: dict, doc_number: str = '1',
                      lang: str = 'auto') -> Document:
    """
    Создаёт альбомный акт списания основных средств.
    group = {
        'inv_number': str,
        'devices': [{'name': str, 'parts': [...]}],
        'org_name': str,
        'region': str,
        'commission_head': str,
        'member1': str,
        'member2': str,
        'doc_date': str,         # дд.мм.гггг
        'narxi': str,            # стоимость (опц.)
        'amort': str,            # амортизация (опц.)
        'norma': str,            # норма (опц.)
        'qoldiq': str,           # остаток (опц.)
        'debet': str,            # счёт дебет (опц.)
        'kredit': str,           # счёт кредит (опц.)
        'bulim': str,            # подразделение (опц.)
        'order_info': str,       # реквизиты приказа (опц.)
        'year_made': str,        # год изготовления (опц.)
        'year_arrived': str,     # год поступления (опц.)
        'year_started': str,     # год ввода (опц.)
    }
    """
    if lang == 'auto':
        lang = _auto_lang(group)

    T = TEXTS_CYR if lang == 'cyrillic' else TEXTS_LAT
    M = MONTHS[lang]

    org_name = convert_to(group.get('org_name', ''),        lang)
    region   = convert_to(group.get('region', ''),          lang)
    boss     = convert_to(group.get('commission_head', ''), lang)
    eng1     = convert_to(group.get('member1', ''),         lang)
    eng2     = convert_to(group.get('member2', ''),         lang)
    inv      = group.get('inv_number', '')

    devices  = group.get('devices', [])

    today    = datetime.now()
    date_str = group.get('doc_date', today.strftime('%d.%m.%Y'))
    try:
        dt = datetime.strptime(date_str, '%d.%m.%Y')
        day = str(dt.day); month = M[dt.month]; year = str(dt.year)
    except Exception:
        day = str(today.day); month = M[today.month]; year = str(today.year)

    # Доп. поля
    narxi      = group.get('narxi', '')
    amort      = group.get('amort', '0.00')
    norma      = group.get('norma', '')
    qoldiq     = group.get('qoldiq', '0.00')
    debet      = group.get('debet', '9210')
    kredit     = group.get('kredit', '0150')
    bulim      = convert_to(group.get('bulim', org_name), lang)
    order_info = group.get('order_info', '')
    year_made  = group.get('year_made', '')
    year_arr   = group.get('year_arrived', '')
    year_start = group.get('year_started', '')

    # Все устройства из группы
    dev_names = [d.get('name','').strip() for d in devices if d.get('name','').strip()]
    names_str = ', '.join(dev_names)

    doc = Document()
    _set_landscape(doc)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # ══════════════════════
    #  ШАПКА — справа
    # ══════════════════════
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after  = Pt(0)

    def rr(p, text, bold=False, size=10):
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size); r.font.name = 'Times New Roman'

    rr(p_hdr, f'{T["tasdiq"]}\n', bold=True, size=10)
    rr(p_hdr, f'{org_name}\n', size=10)
    rr(p_hdr, f'___________ {boss}\n', size=10)
    rr(p_hdr, f'{year} й. «_____»  ____________', size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ══════════════════════
    #  ЗАГОЛОВОК
    # ══════════════════════
    _p(doc, T['title1'], bold=True, size=12,
       align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=0)
    _p(doc, f'№ {doc_number}  - СОН {T["title2"]}',
       bold=True, size=12,
       align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)

    # ══════════════════════
    #  ТАБЛИЦА
    # ══════════════════════
    # Колонки: Бўлим | Дебет | Кредит | Номи | Нархи | Инв.№ | Амортизация | Норма | Қолдиқ
    col_widths = [Cm(3.0), Cm(2.2), Cm(2.2), Cm(6.0), Cm(3.0),
                  Cm(2.8), Cm(3.2), Cm(2.5), Cm(2.8)]
    col_keys   = ['bulim','debet','kredit','nomi','narxi',
                  'inv','amort','norma','qoldiq']
    HDR_FILL   = '2E4057'

    table = doc.add_table(rows=2, cols=len(col_widths))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Строка заголовков
    hdr_row = table.rows[0]
    hdr_row.height = Twips(900)
    for ci, (w, key) in enumerate(zip(col_widths, col_keys)):
        cell = hdr_row.cells[ci]
        cell.width = w
        _set_cell_bg(cell, HDR_FILL)
        _set_cell_border(cell, HDR_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_p(cell, T[key], bold=True, size=8,
                align=WD_ALIGN_PARAGRAPH.CENTER, color=(255,255,255))

    # Строка данных
    data_row = table.rows[1]
    data_row.height = Twips(400)
    row_values = [bulim, debet, kredit, names_str, narxi,
                  inv, amort, norma, qoldiq]
    for ci, (w, val) in enumerate(zip(col_widths, row_values)):
        cell = data_row.cells[ci]
        cell.width = w
        _set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_p(cell, val, size=9,
                align=WD_ALIGN_PARAGRAPH.LEFT if ci in (0,3) else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ══════════════════════
    #  ТЕКСТ ЗАКЛЮЧЕНИЯ
    # ══════════════════════
    # Комиссия
    comm_text = T['komissiya'].format(
        org=org_name,
        date=year,
        order=order_info or '___'
    )
    _p(doc, comm_text, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=2)

    # История оборудования
    for yr, key in [(year_made,'ish_chiq'),(year_arr,'kelgan'),(year_start,'ishga')]:
        _p(doc, f'{yr or "____"} {T[key]}', size=10, before=0, after=1)

    _p(doc, T['kapital'],  size=10, before=0, after=1)
    _p(doc, T['texnik'],   size=10, before=0, after=2)

    # Хулоса
    xulosa = T['xulosa_txt'].format(
        org=org_name,
        names=names_str,
        date=year,
        day=day,
        month=month,
    )
    _p(doc, xulosa, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
       before=0, after=6)

    # ══════════════════════
    #  ПОДПИСИ
    # ══════════════════════
    # Моддий жавобгар
    p_mod = doc.add_paragraph()
    p_mod.paragraph_format.space_before = Pt(0)
    p_mod.paragraph_format.space_after  = Pt(4)
    r1 = p_mod.add_run(f'{T["moddiy"]}  _______________ ')
    r1.font.size = Pt(10); r1.font.name = 'Times New Roman'
    r2 = p_mod.add_run(boss)
    r2.bold = True; r2.font.size = Pt(10); r2.font.name = 'Times New Roman'

    # Комиссия раиси
    p_rais = doc.add_paragraph()
    p_rais.paragraph_format.space_before = Pt(0)
    p_rais.paragraph_format.space_after  = Pt(3)
    r1 = p_rais.add_run(f'{T["komrais"].format(org=org_name)}')
    r1.font.size = Pt(10); r1.font.name = 'Times New Roman'
    r2 = p_rais.add_run(f'                    _____________ {boss}')
    r2.font.size = Pt(10); r2.font.name = 'Times New Roman'

    # Члены комиссии
    p_az = doc.add_paragraph()
    p_az.paragraph_format.space_before = Pt(0)
    p_az.paragraph_format.space_after  = Pt(2)
    rr_az = p_az.add_run(T['komazolari'])
    rr_az.font.size = Pt(10); rr_az.font.name = 'Times New Roman'

    for engineer in [eng1, eng2]:
        if not engineer:
            continue
        p_e = doc.add_paragraph()
        p_e.paragraph_format.space_before = Pt(0)
        p_e.paragraph_format.space_after  = Pt(2)
        r_t = p_e.add_run(f'{T["muhandis"]}                                           ____________')
        r_t.font.size = Pt(10); r_t.font.name = 'Times New Roman'
        r_n = p_e.add_run(engineer)
        r_n.font.size = Pt(10); r_n.font.name = 'Times New Roman'

    return doc



# ─────────────────────────────────────────────
#  Сохранение — каждая группа отдельным файлом
# ─────────────────────────────────────────────

def _safe_filename(text: str) -> str:
    text = text.replace('\n','_').replace('\r','').replace('\t','_')
    for ch in r'\/:*?"<>|':
        text = text.replace(ch, '-')
    return '_'.join(text.split())[:60] or 'doc'


def save_single_files(groups: list, output_dir: str,
                      lang: str = 'auto') -> list:
    """Один .docx на каждую группу (инвентарный номер)."""
    os.makedirs(output_dir, exist_ok=True)
    created = []
    for i, group in enumerate(groups, start=1):
        inv  = _safe_filename(group.get('inv_number', str(i)))
        doc  = create_spisan_doc(group, doc_number=str(i), lang=lang)
        name = f'Akt_Spisaniya_{inv}.docx'
        path = os.path.join(output_dir, name)
        doc.save(path)
        created.append(path)
    return created


# ─────────────────────────────────────────────
#  Сохранение — все группы в одном файле
#  Единая таблица + единое заключение
# ─────────────────────────────────────────────

def save_all_in_one(groups: list, output_path: str,
                    lang: str = 'auto') -> str:
    """
    Все группы в одном альбомном документе.
    Единая шапка, единая таблица (каждая группа — строка),
    единое заключение и подписи.
    """
    if not groups:
        return output_path

    if lang == 'auto':
        lang = _auto_lang(groups[0])

    T  = TEXTS_CYR if lang == 'cyrillic' else TEXTS_LAT
    M  = MONTHS[lang]

    first    = groups[0]
    org_name = convert_to(first.get('org_name', ''),        lang)
    boss     = convert_to(first.get('commission_head', ''), lang)
    eng1     = convert_to(first.get('member1', ''),         lang)
    eng2     = convert_to(first.get('member2', ''),         lang)

    today    = datetime.now()
    date_str = first.get('doc_date', today.strftime('%d.%m.%Y'))
    try:
        dt = datetime.strptime(date_str, '%d.%m.%Y')
        day = str(dt.day); month = M[dt.month]; year = str(dt.year)
    except Exception:
        day = str(today.day); month = M[today.month]; year = str(today.year)

    doc = Document()
    _set_landscape(doc)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # ══ ШАПКА ══
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after  = Pt(0)
    def rr(p, text, bold=False, size=10):
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size); r.font.name = 'Times New Roman'
    rr(p_hdr, f'{T["tasdiq"]}\n', bold=True)
    rr(p_hdr, f'{org_name}\n')
    rr(p_hdr, f'___________ {boss}\n')
    rr(p_hdr, f'{year} й. «_____»  ____________')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ══ ЗАГОЛОВОК ══
    _p(doc, T['title1'], bold=True, size=12,
       align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=0)
    _p(doc, f'№ 1 - СОН {T["title2"]}',
       bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)

    # ══ ЕДИНАЯ ТАБЛИЦА ══
    col_widths = [Cm(3.0), Cm(2.2), Cm(2.2), Cm(5.5), Cm(3.0),
                  Cm(2.8), Cm(3.2), Cm(2.5), Cm(2.8)]
    col_keys   = ['bulim','debet','kredit','nomi','narxi',
                  'inv','amort','norma','qoldiq']
    HDR_FILL   = '2E4057'

    table = doc.add_table(rows=1, cols=len(col_widths))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовок — один раз
    hdr_row = table.rows[0]
    hdr_row.height = Twips(900)
    for ci, (w, key) in enumerate(zip(col_widths, col_keys)):
        cell = hdr_row.cells[ci]
        cell.width = w
        _set_cell_bg(cell, HDR_FILL)
        _set_cell_border(cell, HDR_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_p(cell, T[key], bold=True, size=8,
                align=WD_ALIGN_PARAGRAPH.CENTER, color=(255,255,255))

    all_names = []

    # Строки данных — по одной на каждую группу
    for group in groups:
        inv      = group.get('inv_number', '')
        devices  = group.get('devices', [])
        dev_names = [d.get('name','').strip() for d in devices if d.get('name','').strip()]
        names_str = ', '.join(dev_names)
        if names_str:
            all_names.append(names_str)

        bulim_v  = convert_to(group.get('bulim', org_name), lang)
        narxi    = group.get('narxi', '')
        amort    = group.get('amort', '0.00')
        norma    = group.get('norma', '')
        qoldiq   = group.get('qoldiq', '0.00')
        debet    = group.get('debet', '9210')
        kredit   = group.get('kredit', '0150')

        data_row = table.add_row()
        data_row.height = Twips(400)
        row_values = [bulim_v, debet, kredit, names_str, narxi,
                      inv, amort, norma, qoldiq]
        for ci, (w, val) in enumerate(zip(col_widths, row_values)):
            cell = data_row.cells[ci]
            cell.width = w
            _set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _cell_p(cell, val, size=9,
                    align=WD_ALIGN_PARAGRAPH.LEFT if ci in (0,3)
                    else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ══ ЗАКЛЮЧЕНИЕ (единое) ══
    order_info = first.get('order_info', '___')
    comm_text  = T['komissiya'].format(org=org_name, date=year, order=order_info)
    _p(doc, comm_text, size=10, before=0, after=2)

    year_made  = first.get('year_made', '____')
    year_arr   = first.get('year_arrived', '____')
    year_start = first.get('year_started', '____')
    for yr, key in [(year_made,'ish_chiq'),(year_arr,'kelgan'),(year_start,'ishga')]:
        _p(doc, f'{yr} {T[key]}', size=10, before=0, after=1)
    _p(doc, T['kapital'], size=10, before=0, after=1)
    _p(doc, T['texnik'],  size=10, before=0, after=2)

    all_names_str = '; '.join(all_names)
    xulosa = T['xulosa_txt'].format(
        org=org_name, names=all_names_str,
        date=year, day=day, month=month)
    _p(doc, xulosa, size=10, before=0, after=6)

    # ══ ПОДПИСИ ══
    p_mod = doc.add_paragraph()
    p_mod.paragraph_format.space_before = Pt(0)
    p_mod.paragraph_format.space_after  = Pt(4)
    r1 = p_mod.add_run(f'{T["moddiy"]}  _______________ ')
    r1.font.size = Pt(10); r1.font.name = 'Times New Roman'
    r2 = p_mod.add_run(boss)
    r2.bold = True; r2.font.size = Pt(10); r2.font.name = 'Times New Roman'

    p_rais = doc.add_paragraph()
    p_rais.paragraph_format.space_before = Pt(0)
    p_rais.paragraph_format.space_after  = Pt(3)
    r1 = p_rais.add_run(f'{T["komrais"].format(org=org_name)}')
    r1.font.size = Pt(10); r1.font.name = 'Times New Roman'
    r2 = p_rais.add_run(f'                    _____________ {boss}')
    r2.font.size = Pt(10); r2.font.name = 'Times New Roman'

    p_az = doc.add_paragraph()
    p_az.paragraph_format.space_before = Pt(0)
    p_az.paragraph_format.space_after  = Pt(2)
    rr_az = p_az.add_run(T['komazolari'])
    rr_az.font.size = Pt(10); rr_az.font.name = 'Times New Roman'

    for engineer in [eng1, eng2]:
        if not engineer:
            continue
        p_e = doc.add_paragraph()
        p_e.paragraph_format.space_before = Pt(0)
        p_e.paragraph_format.space_after  = Pt(2)
        r_t = p_e.add_run(f'{T["muhandis"]}                                           ____________')
        r_t.font.size = Pt(10); r_t.font.name = 'Times New Roman'
        r_n = p_e.add_run(engineer)
        r_n.font.size = Pt(10); r_n.font.name = 'Times New Roman'

    doc.save(output_path)
    return output_path
